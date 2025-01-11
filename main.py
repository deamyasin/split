import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox, Scrollbar, ttk
from PIL import Image, ImageTk
import fitz  # PyMuPDF
import os
import threading
import json
from datetime import datetime
from docx import Document


class DeaPDFSplitorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Dea PDF Splitor")
        self.pdf_file = None
        self.page_images = []
        self.current_page = 0
        self.pdf_doc = None
        self.zoom_level = 1.0
        self.split_marks = []  # Store split points
        
        self.thumbnails = []  # Store references to thumbnails for alignment
        
        self.setup_ui()
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
        self.bind_keys()
        self.root.bind("<Shift-Right>", self.rotate_right)
        self.root.bind("<Shift-Left>", self.rotate_left)
        self.root.bind("<Delete>", self.delete_current_page)
        self.root.bind('<Shift-plus>', self.zoom_in)  # Zoom in
        self.root.bind('<Control-minus>', self.zoom_out)  # Zoom out

        self.page_rotation = [0] * 100  # Track rotation up to 100 pages

    def setup_ui(self):
        # Header Frame
        header = tk.Frame(self.root, bg='lightblue', padx=10, pady=10)
        header.pack(side=tk.TOP, fill=tk.X)

        # Current file name label
        self.file_name_label = tk.Label(header, text="", bg='lightblue', font=("Arial", 12))
        self.file_name_label.pack(side=tk.LEFT, padx=10)

        self.setup_buttons(header)

        # Left: Thumbnail Sidebar
        self.thumbnail_frame = tk.Frame(self.root, width=150, bg='white')  # Reduced sidebar width
        self.thumbnail_frame.pack(side=tk.LEFT, fill=tk.Y)

        self.setup_thumbnail_canvas()

        # Right: Marker List
        self.marker_list_frame = tk.Frame(self.root, width=200, bg='lightgrey')
        self.marker_list_frame.pack(side=tk.RIGHT, fill=tk.Y)

        self.marker_list_label = tk.Label(self.marker_list_frame, text="Penanda", bg='lightgrey')
        self.marker_list_label.pack(pady=5)

        self.marker_listbox = tk.Listbox(self.marker_list_frame)
        self.marker_listbox.pack(fill=tk.BOTH, expand=True, pady=10)
        self.marker_listbox.bind("<Double-Button-1>", self.on_marker_double_click)

        self.remove_marker_button = tk.Button(self.marker_list_frame, text="Hapus Penanda", state="disabled", command=self.remove_split_mark)
        self.remove_marker_button.pack(pady=10)

        # Add horizontal scrollbar
        marker_scrollbar = Scrollbar(self.marker_list_frame, orient="horizontal", command=self.marker_listbox.xview)
        marker_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        self.marker_listbox.config(xscrollcommand=marker_scrollbar.set)

        self.marker_listbox.bind("<Double-Button-1>", self.on_marker_double_click)

        # Center: PDF Preview Area
        self.preview_frame = tk.Frame(self.root, bg='grey', width=500, height=600)
        self.preview_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        self.setup_preview_canvas()

        # Pagination Label
        self.pagination_label = tk.Label(self.preview_frame, text="Halaman - dari -", bg="lightgrey", font=("Arial", 10))
        self.pagination_label.pack(side=tk.BOTTOM, pady=5)


        # Entry untuk input halaman
        self.page_entry = tk.Entry(self.preview_frame, width=5)
        self.page_entry.bind("<space>", self.go_to_page)
        self.page_entry.pack(side=tk.BOTTOM)

    def on_thumbnail_frame_configure(self, event):
        self.thumbnail_canvas.configure(scrollregion=self.thumbnail_canvas.bbox("all"))
    
    def scroll_to_active_thumbnail(self):
        for idx, widget in enumerate(self.thumbnail_inner_frame.winfo_children()):
            if isinstance(widget, tk.Button) and idx == self.current_page:
                y = widget.winfo_y()
                self.thumbnail_canvas.yview_moveto(y / self.thumbnail_inner_frame.winfo_height())
                break
    
    def setup_buttons(self, parent):
        buttons = [
            ("Pick File", self.open_file),
            ("Export Marker", self.save_markers_to_file),
            ("Load Marker", self.load_markers_from_file),
            ("Zoom In (Shift +)", self.zoom_in),
            ("Zoom Out (Ctrl -)", self.zoom_out),
            ("Add Split PDF (ENTER)", self.add_split_mark, "orange"),
            ("Rotate Left (Shift Kiri)", self.rotate_left),
            ("Rotate Right (Shift Kanan)", self.rotate_right),
            ("Proses Output", self.output_splits, "green"),
            ("Convert To Word", self.convert_to_docx, "blue"),
        ]

        for (text, command, *color) in buttons:
            btn = tk.Button(parent, text=text, command=command, bg=color[0] if color else None)
            btn.pack(side=tk.LEFT, padx=5)

        # Store reference to "Add Split PDF" button
        self.add_marker_button = btn  # "Add Split PDF" button reference
    
    def rotate_left(self, event=None):
        self.page_rotation[self.current_page] = (self.page_rotation[self.current_page] - 90) % 360
        self.load_page_preview(self.current_page)

    def rotate_right(self, event=None):
        self.page_rotation[self.current_page] = (self.page_rotation[self.current_page] + 90) % 360
        self.load_page_preview(self.current_page)

    # def setup_thumbnail_canvas(self):
    #     self.thumbnail_canvas = tk.Canvas(self.thumbnail_frame)
    #     self.thumbnail_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
    #     self.thumbnail_scrollbar = Scrollbar(self.thumbnail_frame, orient="vertical", command=self.thumbnail_canvas.yview)
    #     self.thumbnail_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    #     self.thumbnail_canvas.configure(yscrollcommand=self.thumbnail_scrollbar.set)
        
    #     self.thumbnail_inner_frame = tk.Frame(self.thumbnail_canvas)
    #     self.thumbnail_canvas.create_window((0, 0), window=self.thumbnail_inner_frame, anchor='nw')

    #     self.thumbnail_inner_frame.bind("<Configure>", self.on_thumbnail_frame_configure)

    def setup_preview_canvas(self):
        self.preview_canvas = tk.Canvas(self.preview_frame, bg="#EEEEFF")
        self.preview_canvas.pack(fill=tk.BOTH, expand=True)
        self.preview_image_id = None

    def open_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if file_path:
            threading.Thread(target=self.load_pdf, args=(file_path,)).start()

    def load_pdf(self, file_path):
        try:
            self.pdf_file = file_path
            self.file_name_label.config(text=os.path.basename(file_path))
            self.pdf_doc = fitz.open(file_path)
            self.current_page = 0        
            self.split_marks.clear()
            self.page_rotation = [0] * len(self.pdf_doc)
            self.load_thumbnails()
            self.load_page_preview(self.current_page)
        except Exception as e:
            messagebox.showerror("Error", f"Could not load PDF file: {e}")
    
    def add_split_mark(self):
        if self.current_page_has_marker():
            messagebox.showwarning("Warning", "Halaman ini sudah memiliki penanda!")
            return

        split_name = simpledialog.askstring("Input", "Masukkan nama untuk tanda split:")
        if split_name:
            self.split_marks.append((self.current_page, split_name))
            self.split_marks.sort()  # Automatically sort by page number
            self.update_marker_list()
            self.update_thumbnail_borders()


    def update_marker_list(self):
        self.marker_listbox.delete(0, tk.END)
        for page_num, split_name in self.split_marks:
            self.marker_listbox.insert(tk.END, f"Halaman {page_num + 1}: {split_name}")

        # Disable add marker button if current page has marker
        self.update_marker_button_states()

    def update_marker_button_states(self):
        # Check if current page has a marker
        has_marker = self.current_page_has_marker()
        self.add_marker_button.config(state="disabled" if has_marker else "normal")
        self.remove_marker_button.config(state="normal" if has_marker else "disabled")

    def current_page_has_marker(self):
        return any(mark[0] == self.current_page for mark in self.split_marks)

    def on_marker_double_click(self, event):
        selection = self.marker_listbox.curselection()
        if selection:
            index = selection[0]
            page_num, split_name = self.split_marks[index]
            new_name = simpledialog.askstring("Rename", "Masukkan nama baru:", initialvalue=split_name)
            if new_name:
                self.split_marks[index] = (page_num, new_name)
                self.update_marker_list()

    def remove_split_mark(self):
        self.split_marks = [mark for mark in self.split_marks if mark[0] != self.current_page]
        self.update_marker_list()
        self.update_thumbnail_borders()

    def load_page_preview(self, page_num):
        self.current_page = page_num
        page = self.pdf_doc.load_page(page_num)
        rotation_angle = self.page_rotation[self.current_page]
        page.set_rotation(rotation_angle)
        zoom_matrix = fitz.Matrix(self.zoom_level, self.zoom_level)
        pix = page.get_pixmap(matrix=zoom_matrix)

        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        img_tk = ImageTk.PhotoImage(img)
        
        self.preview_canvas.delete("all")
        self.preview_image_id = self.preview_canvas.create_image(0, 0, anchor='nw', image=img_tk)
        self.preview_canvas.image = img_tk

        self.update_pagination()
        self.update_thumbnail_borders()

        # Update button states after loading new page
        self.update_marker_button_states()        
        self.scroll_to_active_thumbnail()

    def load_thumbnails(self):
        for widget in self.thumbnail_inner_frame.winfo_children():
            widget.destroy()

        for page_num in range(len(self.pdf_doc)):
            self.create_thumbnail(page_num)

    def create_thumbnail(self, page_num):
        page = self.pdf_doc.load_page(page_num)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        img.thumbnail((120, 120))  # Adjust thumbnail size
        img_tk = ImageTk.PhotoImage(img)

        thumbnail_btn = tk.Button(self.thumbnail_inner_frame, image=img_tk, command=lambda p=page_num: self.load_page_preview(p))
        thumbnail_btn.image = img_tk
        thumbnail_btn.pack(pady=5)

        if page_num == self.current_page:
            thumbnail_btn.config(bg="red")
        else:
            thumbnail_btn.config(bg="white")

    def update_thumbnail_borders(self):
        for idx, widget in enumerate(self.thumbnail_inner_frame.winfo_children()):
            if isinstance(widget, tk.Button):
                widget.config(bg="red" if idx == self.current_page else "white")

    # def update_pagination(self):
    #     total_pages = len(self.pdf_doc)
    #     self.pagination_label.config(text=f"Halaman {self.current_page + 1} dari {total_pages}")
    
    def update_pagination(self):
        """Memperbarui label pagination dan menampilkan halaman saat ini."""
        total_pages = len(self.pdf_doc)
        self.pagination_label.config(text=f"Halaman {self.current_page + 1} dari {total_pages}")
        self.page_entry.delete(0, tk.END)
        self.page_entry.insert(0, str(self.current_page + 1))

    def zoom_in(self, event=None):
        self.zoom_level += 0.2
        self.load_page_preview(self.current_page)

    def zoom_out(self, event=None):
        self.zoom_level -= 0.1
        self.load_page_preview(self.current_page)

    def output_splits(self):
        output_folder = filedialog.askdirectory(title="Pilih Folder Output")
        if output_folder:
            threading.Thread(target=self.process_split_output, args=(output_folder,)).start()
    
    def process_split_output(self, output_folder):
        if not self.split_marks:
            messagebox.showinfo("Info", "Tidak ada penanda yang ditambahkan!")
            return

        # Ensure markers are sorted by page number
        self.split_marks.sort()

        # Ensure split marks are within valid range and no duplicates
        unique_marks = sorted(set(self.split_marks), key=lambda x: x[0])

        for i, (page_num, split_name) in enumerate(unique_marks):
            if i == len(unique_marks) - 1:
                next_page_num = len(self.pdf_doc)
            else:
                next_page_num = unique_marks[i + 1][0]

            output_pdf = fitz.open()
            for page_index in range(page_num, next_page_num):
                output_pdf.insert_pdf(self.pdf_doc, from_page=page_index, to_page=page_index)

            output_filename = f"{split_name}.pdf"
            output_path = os.path.join(output_folder, output_filename)
            output_pdf.save(output_path)
            output_pdf.close()

        messagebox.showinfo("Success", "Splits PDF berhasil disimpan!")


    def delete_current_page(self, event=None):
        if messagebox.askyesno("Confirm Deletion", "Are you sure you want to delete this page?"):
            try:
                self.pdf_doc.delete_page(self.current_page)
                del self.page_rotation[self.current_page]  # Remove the rotation for the deleted page
                if self.current_page >= len(self.pdf_doc):
                    self.current_page = max(0, len(self.pdf_doc) - 1)  # Adjust current page
                self.load_thumbnails()
                self.load_page_preview(self.current_page)  # Load new current page
                messagebox.showinfo("Success", "Page deleted successfully.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to delete page: {e}")
    
    def save_markers_to_file(self):
        if not self.split_marks:
            messagebox.showinfo("Info", "Tidak ada penanda untuk disimpan!")
            return

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        markers_filename = f"{os.path.basename(self.pdf_file).split('.')[0]}_{timestamp}_mark.txt"
        markers_path = os.path.join(os.path.dirname(self.pdf_file), markers_filename)

        with open(markers_path, 'w') as file:
            json.dump(self.split_marks, file)
        
        messagebox.showinfo("Success", f"Penanda disimpan di {markers_path}")

    def load_markers_from_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Marker files", "*.txt")])
        if file_path:
            with open(file_path, 'r') as file:
                self.split_marks = json.load(file)
            self.update_marker_list()
            self.update_thumbnail_borders()
            messagebox.showinfo("Success", "Penanda berhasil dimuat!")


    def go_to_page(self, event):
        """Berpindah ke halaman yang diinput dalam kotak Entry."""
        total_pages = len(self.pdf_doc)
        try:
            page_num = int(self.page_entry.get()) - 1  # Konversi input ke index halaman (mulai dari 0)
            lambda e: self.load_page_preview(page_num)
            if 0 <= page_num < total_pages:
                self.current_page = page_num
                self.update_pagination()
            else:
                self.page_entry.delete(0, tk.END)
                self.page_entry.insert(0, str(self.current_page + 1))  # Reset jika di luar range
        except ValueError:
            # Jika input bukan angka, kembalikan ke halaman sekarang
            self.page_entry.delete(0, tk.END)
            self.page_entry.insert(0, str(self.current_page + 1))

    def bind_keys(self):
        self.root.bind("<Left>", lambda e: self.load_page_preview(max(0, self.current_page - 1)))
        self.root.bind("<Right>", lambda e: self.load_page_preview(min(len(self.pdf_doc) - 1, self.current_page + 1)))
        self.root.bind("<Up>", lambda e: self.load_page_preview(max(0, self.current_page - 1)))
        self.root.bind("<Down>", lambda e: self.load_page_preview(min(len(self.pdf_doc) - 1, self.current_page + 1)))
        self.root.bind("<Return>", lambda e: self.add_split_mark())
        # self.root.bind('<Alt-Next>', self.move_page_down)  # Alt + Page Down
        # self.root.bind('<Alt-Prior>', self.move_page_up)  # Alt + Page Up

    def parse_page_selection(self, input_str, total_pages):
        """Fungsi untuk mengubah input halaman (misal: 1,3-5,7) menjadi daftar halaman."""
        selected_pages = set()
        for part in input_str.split(','):
            if '-' in part:
                start, end = map(int, part.split('-'))
                selected_pages.update(range(start, end + 1))
            else:
                selected_pages.add(int(part))
        
        # Validasi halaman yang dipilih harus berada dalam range yang tersedia
        selected_pages = sorted([p for p in selected_pages if 1 <= p <= total_pages])
        return selected_pages

    def convert_to_docx(self):
        """Fungsi untuk memilih halaman dan mengonversi ke DOCX."""
        if not self.pdf_doc:
            messagebox.showwarning("Warning", "No PDF loaded!")
            return

        # Total halaman tersedia di PDF
        total_pages = len(self.pdf_doc)

        # Input halaman dari pengguna
        page_range = simpledialog.askstring(
            "Convert to DOCX", f"Enter pages to convert (1-{total_pages}), e.g., 1,3-5:"
        )

        if page_range:
            try:
                selected_pages = self.parse_page_selection(page_range, total_pages)
                if not selected_pages:
                    raise ValueError

                # Memilih lokasi untuk menyimpan file DOCX
                output_file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
                if not output_file:
                    return

                # Proses konversi PDF ke DOCX
                self.create_docx(selected_pages, output_file)
                messagebox.showinfo("Success", f"File successfully converted to {output_file}")

            except ValueError:
                messagebox.showerror("Error", "Invalid page selection input!")
    
    def create_docx(self, selected_pages, output_file):
        """Fungsi untuk mengonversi halaman yang dipilih ke DOCX."""
        doc = Document()

        for page_number in selected_pages:
            page = self.pdf_doc.load_page(page_number - 1)  # MuPDF 0-indexed, halaman user 1-indexed
            text = page.get_text("text")
            doc.add_paragraph(f"Page {page_number}")
            doc.add_paragraph(text)
            doc.add_page_break()

        # Simpan ke file DOCX
        doc.save(output_file)
    
    def setup_thumbnail_canvas(self):
        self.thumbnail_canvas = tk.Canvas(self.thumbnail_frame)
        self.thumbnail_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.thumbnail_scrollbar = Scrollbar(self.thumbnail_frame, orient="vertical", command=self.thumbnail_canvas.yview)
        self.thumbnail_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.thumbnail_canvas.configure(yscrollcommand=self.thumbnail_scrollbar.set)
        
        self.thumbnail_inner_frame = tk.Frame(self.thumbnail_canvas)
        self.thumbnail_canvas.create_window((0, 0), window=self.thumbnail_inner_frame, anchor='nw')

        self.thumbnail_inner_frame.bind("<Configure>", self.on_thumbnail_frame_configure)

    # def create_thumbnail(self, page_num):
    #     page = self.pdf_doc.load_page(page_num)
    #     pix = page.get_pixmap()
    #     img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    #     img.thumbnail((120, 120))
    #     img_tk = ImageTk.PhotoImage(img)

    #     thumbnail_btn = tk.Button(self.thumbnail_inner_frame, image=img_tk)
    #     thumbnail_btn.image = img_tk
    #     thumbnail_btn.pack(pady=5)

    #     # # Binding drag and drop events for thumbnail
    #     # thumbnail_btn.bind("<ButtonPress-1>", lambda e, idx=page_num: self.start_drag(e, idx))
    #     # thumbnail_btn.bind("<B1-Motion>", self.on_drag)
    #     # thumbnail_btn.bind("<ButtonRelease-1>", self.end_drag)

    #     self.thumbnails.append(thumbnail_btn)

    #     if page_num == self.current_page:
    #         thumbnail_btn.config(bg="red")
    #     else:
    #         thumbnail_btn.config(bg="white")

    # def start_drag(self, event, index):
    #     self.dragging_index = index
    #     self.placeholder = tk.Frame(self.thumbnail_inner_frame, width=event.widget.winfo_width(), height=event.widget.winfo_height(), bg="grey")
    #     self.placeholder.pack_forget()

    # def on_drag(self, event):
    #     if self.dragging_index is not None:
    #         widget = event.widget
    #         widget.lift()  # Bring dragged widget to front
    #         widget.place(x=event.widget.winfo_x(), y=event.y_root - widget.winfo_rooty())

    #         # Calculate the nearest position for placeholder
    #         drop_y = event.y_root - self.thumbnail_inner_frame.winfo_rooty()
    #         new_index = self.get_drop_index(drop_y)

    #         # Move the placeholder if position changes
    #         if self.placeholder.winfo_ismapped():
    #             self.placeholder.pack_forget()
    #         self.placeholder.pack(pady=5, before=self.thumbnails[new_index])

    # def end_drag(self, event):
    #     if self.dragging_index is not None:
    #         widget = event.widget
    #         widget.place_forget()  # Stop manual placement
    #         drop_y = event.y_root - self.thumbnail_inner_frame.winfo_rooty()

    #         new_index = self.get_drop_index(drop_y)

    #         # Reorder thumbnails and pages
    #         self.reorder_thumbnails(self.dragging_index, new_index)
    #         self.update_preview_and_markers(new_index)

    #         self.dragging_index = None
    #         if self.placeholder.winfo_ismapped():
    #             self.placeholder.pack_forget()

    # def get_drop_index(self, drop_y):
    #     # Determine the index where the thumbnail is dropped based on the y position
    #     for i, widget in enumerate(self.thumbnail_inner_frame.winfo_children()):
    #         widget_y = widget.winfo_y()
    #         if drop_y < widget_y + widget.winfo_height() // 2:
    #             return i
    #     return len(self.thumbnails) - 1  # If dropped below the last thumbnail

    # def reorder_thumbnails(self, from_index, to_index):
    #     # Swap the thumbnail order in the UI
    #     if from_index != to_index:
    #         # Move the thumbnails in the list
    #         self.thumbnails[from_index], self.thumbnails[to_index] = self.thumbnails[to_index], self.thumbnails[from_index]

    #         # Rearrange the thumbnail buttons in the UI
    #         self.thumbnails[from_index].pack_forget()
    #         self.thumbnails[to_index].pack_forget()
    #         self.thumbnails[to_index].pack(pady=5)
    #         self.thumbnails[from_index].pack(pady=5)

    #         # Reorder PDF pages internally (example using PyMuPDF or similar)
    #         self.pdf_doc.insert_pdf(self.pdf_doc, from_page=from_index, to_page=from_index)
    #         self.pdf_doc.delete_page(from_index if from_index > to_index else from_index + 1)

    # def update_preview_and_markers(self, new_index):
    #     # Automatically update the PDF preview to the new index page
    #     self.current_page = new_index
    #     self.update_pdf_preview()

    #     # If the page has a marker, update the marker to its new location
    #     if new_index in self.split_marks:
    #         self.update_split_marks(new_index)
    
    # def move_page_down(self, event):
    #     if self.current_page < len(self.page_images) - 1:
    #         # Swap dengan halaman berikutnya
    #         self.page_images[self.current_page], self.page_images[self.current_page + 1] = (
    #             self.page_images[self.current_page + 1],
    #             self.page_images[self.current_page],
    #         )
    #         self.current_page += 1
    #         self.display_page(self.current_page)

    # def move_page_up(self, event):
    #     if self.current_page > 0:
    #         # Swap dengan halaman sebelumnya
    #         self.page_images[self.current_page], self.page_images[self.current_page - 1] = (
    #             self.page_images[self.current_page - 1],
    #             self.page_images[self.current_page],
    #         )
    #         self.current_page -= 1
    #         self.display_page(self.current_page)


    def on_close(self):
        if messagebox.askokcancel("Quit", "Do you want to quit?"):
            self.root.quit()


if __name__ == "__main__":
    root = tk.Tk()
    app = DeaPDFSplitorApp(root)
    root.mainloop()
